import time
import subprocess
from docx import Document
import random
import os
from dotenv import load_dotenv
from Evento import Evento
load_dotenv()

class Proceso:
    contador = 0

    # Lista de comandos permitidos
    comandos_permitidos = {
        "0": "ls -la",
        "1": "dir",
        "2": "ping -c 4 google.com",
        "3":"Systeminfo",
        "4": "Tasklist"
    }

    def __init__(self, prioridad, proceso, tipo_proceso,gestor,tipo_planificacion):
        Proceso.contador += 1
        self.prioridad = prioridad
        self.proceso = proceso
        self.tipo_proceso = tipo_proceso
        self.numero_proceso = Proceso.contador
        self.comandos_permitidos=Proceso.comandos_permitidos
        self.tipo_planificacion=tipo_planificacion
        self.rafaga_cpu = self.calcular_tiempo_ejecucion()
        #Usando de manera adecuada
        self.evento=Evento(gestor)
        self.evento.setEstadoNuevo()
    
    #Obtengo el evento para ser usado en GestorProceso
    def getProceso(self):
        return self.proceso
    def getEvento(self):
        return self.evento
    #Método para Round Robin
    def ha_fallado(self):
        evento=self.evento
        if evento.getEstado()=="Bloqueado":
            return True
        else:
            return False
    def startProcess(self):
        """"Obtengo hora formateada.
        """
        self.tiempo_llegada=time.time()
        self.tiempo_llegada = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(self.tiempo_llegada))

    #Acceder al método sin una instancia
    @classmethod
    def devolver_comandos(self):
        return self.comandos_permitidos
    
    def ejecutar_comando(self):
        """Corre un proceso de una lista, hace un reporte acorde a los resultados.
        """
        try:
            # Verificar si el tipo_proceso está en la lista de comandos permitidos
            if self.tipo_proceso in Proceso.comandos_permitidos:
                comando = Proceso.comandos_permitidos[self.tipo_proceso]
            else:
                raise ValueError(f"Comando no permitido: {self.tipo_proceso}")

            resultado = subprocess.run(comando, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if resultado.returncode == 0:
                salida = resultado.stdout
            else:
                salida = resultado.stderr

            ruta = os.getenv("REPORTS_PATH")
            nombre = f"{ruta}resultado{self.numero_proceso}.docx"
            documento = Document()
            documento.add_heading("Resultado del Proceso", 0)
            documento.add_paragraph(salida)
            documento.save(nombre)
            print("Documento guardado con éxito.")
        except Exception as e:
            print(f"Error al ejecutar el comando: {e}")

    def calcular_tiempo_ejecucion(self):
        """Simula el tiempo de ejecución acorde a la prioridad.
        """
        if self.prioridad == "alta":
            tiempo = random.uniform(1, 3)
        elif self.prioridad == "media":
            tiempo = random.uniform(3, 6)
        elif self.prioridad == "baja":
            tiempo = random.uniform(6, 10)
        else:
            tiempo = 5  
        return tiempo
    
    def ejecutar(self):
        """Ejecuta y pasa filtros acorde a los eventos."""
        try:
           
            self.startProcess()
            # Calcula el tiempo de simulación basado en la prioridad
            
            print(f"Ejecutando proceso con prioridad {self.prioridad} tipo: ({self.tipo_proceso}) ")
            # Simula el tiempo de ejecución
            time.sleep(self.rafaga_cpu)
            print(f"Tiempo estimado del proceso {self.proceso}: {self.rafaga_cpu:.2f} segundos")
            print(f"Momento en que se inicio: {self.tiempo_llegada}")
        except Exception as e:
            print(f"Error al ejecutar el proceso: {e}")
