import time
import subprocess
from docx import Document
import random
import os
from dotenv import load_dotenv
from proceso.Evento import Evento
from output.Reportes import Reportes

load_dotenv()

class Proceso:
    contador = 0

    # Lista de comandos permitidos
    comandos_permitidos = {
        "0": "ls -la",
        "1": "dir",
        "2": "ping -c 4 google.com",
        "3": "Systeminfo",
        "4": "Tasklist"
    }

    def __init__(self, prioridad, proceso, tipo_proceso, gestor, tipo_planificacion):
        Proceso.contador += 1
        self.prioridad = prioridad
        self.proceso = proceso
        self.tipo_proceso = tipo_proceso
        self.numero_proceso = Proceso.contador
        self.comandos_permitidos = Proceso.comandos_permitidos
        self.tipo_planificacion = tipo_planificacion
        self.rafaga_cpu = self.calcular_tiempo_ejecucion()
        # Usando de manera adecuada
        self.evento = Evento(gestor,proceso)
        self.evento.setEstadoNuevo()
        #Usado para los logs
        self.gestor=gestor
    
    # Obtener el evento para ser usado en GestorProceso
    def getProceso(self):
        return self.proceso
    
    def getEvento(self):
        return self.evento

    # Método para Round Robin
    def ha_fallado(self):
        evento = self.evento
        return evento.getEstado() == "Bloqueado"

    def startProcess(self):
        """Obtengo hora formateada."""
        self.tiempo_llegada = time.time()
        self.tiempo_llegada = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(self.tiempo_llegada))

    @classmethod
    def devolver_comandos(cls):
        return cls.comandos_permitidos
    def infLog(self,message):
        log=self.gestor.getLog()
        log.log_info(message)
    def warnLog(self,message):
        log=self.gestor.getLog()
        log.log_warning(message)
    def errLog(self,message):
        log=self.gestor.getLog()
        log.log_error(message)
    def crear_docx(self, nombre, salida):
        try:
            documento = Document()
            documento.add_heading("Resultado del Proceso", 0)
            documento.add_paragraph(salida)
            documento.save(nombre)
        except Exception as e:
            message=f"Error al crear el doc: {e}"
            self.errLog(message)

    def convertir_pdf(self, file):
        reportes = Reportes()
        try:
            reportes.convertir_docx_a_pdf(file,self.gestor)
            message=f"Archivos {file} convertidos a PDF con exito."
            self.infLog(message)
        except Exception as e:
            message=f"Error al convertir los archivos a PDF: {e}"
            self.errLog(message)

    def ejecutar_comando(self):
        """Corre un proceso de una lista, hace un reporte acorde a los resultados."""
        try:
            # Verificar si el tipo_proceso está en la lista de comandos permitidos
            if self.tipo_proceso in Proceso.comandos_permitidos:
                comando = Proceso.comandos_permitidos[self.tipo_proceso]
            else:
                message=f"Comando no permitido: {self.tipo_proceso}"
                self.warnLog(message)
                raise ValueError(f"Comando no permitido: {self.tipo_proceso}")

            resultado = subprocess.run(comando, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            salida = resultado.stdout if resultado.returncode == 0 else resultado.stderr

            ruta = os.getenv("REPORTS_PATH_DOC")
            nombre = f"{ruta}resultado{self.numero_proceso}.docx"
            self.crear_docx(nombre, salida)

            # Convertir el DOCX y las imágenes a PDF
            self.convertir_pdf(file=nombre)
            message="Documentos y PDFS creados con exito."
            self.infLog(message)
        except Exception as e:
            message=f"Error al ejecutar el comando: {e}"
            self.errLog(message)

    def calcular_tiempo_ejecucion(self):
        """Simula el tiempo de ejecución acorde a la prioridad."""
        if self.prioridad == "alta":
            return random.uniform(1, 3)
        elif self.prioridad == "media":
            return random.uniform(3, 6)
        elif self.prioridad == "baja":
            return random.uniform(6, 10)
        return 5 

    def ejecutar(self):
        """Ejecuta y pasa filtros acorde a los eventos."""
        try:
            self.startProcess()
            message=f"Ejecutando proceso con prioridad {self.prioridad}, tipo: ({self.tipo_proceso})"
            self.infLog(message)
            
            # Simula el tiempo de ejecución
            time.sleep(self.rafaga_cpu)
            message=f"Tiempo estimado del proceso {self.proceso}: {self.rafaga_cpu:.2f} segundos"
            self.infLog(message)
            message=f"Momento en que se inicio: {self.tiempo_llegada}"
            self.infLog(message)
        except Exception as e:
            message=f"Error al ejecutar el proceso: {e}"
            self.errLog(message)